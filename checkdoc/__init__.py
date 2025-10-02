"""Scan a word document for key patterns."""

from __future__ import annotations

# Python imports
import argparse
import re
import sys
from typing import Any

# Module imports
from docx import Document


def find_patterns_in_docx(file_path: str, patterns: list[str]) -> list[str]:
    """Find patterns in a docx file."""
    doc = Document(file_path)
    matches = []

    compiled_patterns = [re.compile(p, re.IGNORECASE) for p in patterns]

    def scan_para(obj: Any) -> list[str]:
        matches = []
        for para in obj.paragraphs:
            text = para.text.strip()
            if any(p.search(text) for p in compiled_patterns):
                matches.append(text)
        return matches

    matches += scan_para(doc)

    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                matches += scan_para(cell)

    return matches

def main() -> None:
    """Check the word document from the CLI arguments."""
    parser = argparse.ArgumentParser(
        description="Scan a Word document for key patterns.",
    )
    parser.add_argument("file", help="Path to the .docx file")
    parser.add_argument("patterns", nargs="+", help="Key patterns to search for")
    args = parser.parse_args()

    try:
        results = find_patterns_in_docx(args.file, args.patterns)
    except re.rerror:
        print(f"Invalid regex pattern: {args.patterns}")  # noqa: T201
        sys.exit(1)

    for line in results:
        print(line)  # noqa: T201

    sys.exit(0)

if __name__ == "__main__":
    main()
